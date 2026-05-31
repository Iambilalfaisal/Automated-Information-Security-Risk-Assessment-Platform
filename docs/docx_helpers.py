"""
docx_helpers.py — Low-level python-docx formatting helpers for the IEEE report.

Provides base styling (Times New Roman 12pt, 14pt headings, 1.5 spacing, 1in
margins), bottom-center page numbers starting at the Introduction, a Word TOC
field, numbered/captioned figures and tables, right-aligned numbered equations,
and grey-background Courier New code snippets.
"""

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"


def set_base_styles(doc) -> None:
    """Apply Times New Roman 12pt body, 1.5 spacing, 1in margins, and heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # Restyle built-in heading styles so the TOC field captures them while
    # keeping Times New Roman, bold, black at 14/13/12pt.
    heading_sizes = {"Heading 1": 14, "Heading 2": 13, "Heading 3": 12}
    for name, size in heading_sizes.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_section_heading(doc, text, level=1):
    """Add a heading using a built-in Heading style (captured by the TOC)."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    size = {1: 14, 2: 13, 3: 12}.get(level, 12)
    _set_run_font(run, BODY_FONT, size, bold=True, color=RGBColor(0, 0, 0))
    return heading


def _set_run_font(run, name=BODY_FONT, size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure east-asian font mapping also set (python-docx quirk)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_heading(doc, text, size=14, bold=True, space_before=12):
    """Add a section heading (default 14pt bold)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, size, bold=bold)
    return p


def add_body(doc, text, justify=True, italic=False):
    """Add a justified body paragraph."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, 12, italic=italic)
    return p


def add_bullets(doc, items):
    """Add a bulleted list."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, BODY_FONT, 12)


def add_numbered(doc, items):
    """Add a numbered list."""
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        _set_run_font(run, BODY_FONT, 12)


def add_figure(doc, image_path, fig_no, caption, width_inches=5.5):
    """Insert a centered image with a numbered caption BELOW it."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    crun = cap.add_run(f"Fig. {fig_no}. {caption}")
    _set_run_font(crun, BODY_FONT, 10, italic=True)


def add_table_caption(doc, table_no, caption):
    """Add a numbered table caption ABOVE a table."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    crun = cap.add_run(f"Table {table_no}. {caption}")
    _set_run_font(crun, BODY_FONT, 10, italic=True)


def add_table(doc, headers, rows, col_widths=None):
    """Add a grid table with a header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run_font(run, BODY_FONT, 10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_run_font(run, BODY_FONT, 10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def add_equation(doc, equation_text, number):
    """Add an equation line with a right-aligned equation number."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    # Tab stop at right margin to push the number right
    from docx.enum.text import WD_TAB_ALIGNMENT

    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(equation_text)
    _set_run_font(run, BODY_FONT, 12, italic=True)
    run2 = p.add_run(f"\t({number})")
    _set_run_font(run2, BODY_FONT, 12)


def _shade_paragraph(paragraph, fill="EDEDED"):
    """Apply a grey background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code(doc, code, with_line_numbers=True):
    """Add a code snippet in Courier New 10pt on a grey background."""
    lines = code.strip("\n").split("\n")
    for idx, line in enumerate(lines, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        _shade_paragraph(p)
        prefix = f"{idx:>2}  " if with_line_numbers else ""
        run = p.add_run(prefix + line)
        _set_run_font(run, CODE_FONT, 10, color=RGBColor(0x1A, 0x1A, 0x1A))


def add_page_field(paragraph):
    """Insert a Word PAGE field into a paragraph (auto page number)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)
    _set_run_font(run, BODY_FONT, 11)


def start_numbered_section(doc, restart_at=1):
    """
    Start a new section (page break) whose footer shows a bottom-center page
    number restarting at `restart_at`. Returns the new section.
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    # Restart page numbering
    sectPr = new_section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), str(restart_at))
    sectPr.append(pgNumType)
    # Unlink footer from previous so front matter stays number-free
    new_section.footer.is_linked_to_previous = False
    footer_p = new_section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_p)
    return new_section


def add_toc(doc):
    """Insert an auto-updating Table of Contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose 'Update Field' to populate the Table of Contents."
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(placeholder)
    run._r.append(fldEnd)


def page_break(doc):
    """Insert a manual page break."""
    doc.add_page_break()
