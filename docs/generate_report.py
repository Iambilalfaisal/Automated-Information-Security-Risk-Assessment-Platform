"""
generate_report.py — Generate the IEEE-format technical report as a .docx.

Produces docs/Technical_Report.docx with all 12 mandatory sections, Times New
Roman 12pt body / 14pt bold headings, 1.5 spacing, 1in margins, bottom-center
page numbers starting at the Introduction, numbered/captioned figures and
tables, numbered equations, code snippets, an auto TOC, and IEEE references.

Live figures and result numbers come from a real assessment run on the seeded
demo inventory (see figures.generate_all and seed_data).

Run: python docs/generate_report.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import figures as figmod
from docx_helpers import (
    add_body,
    add_bullets,
    add_code,
    add_equation,
    add_figure,
    add_numbered,
    add_section_heading,
    add_table,
    add_table_caption,
    add_toc,
    page_break,
    set_base_styles,
    start_numbered_section,
    BODY_FONT,
)

OUT_PATH = Path(__file__).resolve().parent / "Technical_Report.docx"
GITHUB_URL = "https://github.com/<your-username>/risk-assessment-platform"


def _centered(doc, text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


# ------------------------------- Sections ---------------------------------


def cover_page(doc):
    """Section 1: Cover page (placeholders for university-specific fields)."""
    for _ in range(2):
        doc.add_paragraph()
    _centered(doc, "[University Logo]", 12, italic=True, space_after=24)
    _centered(doc, "UNIVERSITY OF MANAGEMENT AND TECHNOLOGY", 14, bold=True, space_after=4)
    _centered(doc, "Department of Artificial Intelligence", 12, space_after=36)
    _centered(
        doc,
        "Automated Information Security Risk Assessment Platform",
        18,
        bold=True,
        space_after=8,
    )
    _centered(
        doc,
        "A NIST SP 800-30 and AssessITS Based Quantitative Risk Analysis System",
        12,
        italic=True,
        space_after=48,
    )
    _centered(doc, "Team Members", 12, bold=True, space_after=4)
    _centered(doc, "[Student Name 1] — [Roll Number]", 12, space_after=2)
    _centered(doc, "[Student Name 2] — [Roll Number]", 12, space_after=2)
    _centered(doc, "[Student Name 3] — [Roll Number]", 12, space_after=36)
    _centered(doc, "Course: Information Security  [Course Code]", 12, space_after=2)
    _centered(doc, "Supervisor: [Supervisor Name]", 12, space_after=2)
    _centered(doc, f"Date: {date.today().strftime('%B %d, %Y')}", 12, space_after=2)
    page_break(doc)


def abstract(doc):
    """Section 2: Abstract (150-200 words, no citations/equations)."""
    add_section_heading(doc, "Abstract", level=1)
    text = (
        "Organisations face an expanding landscape of information security threats, "
        "yet many small and medium enterprises lack the tools to quantify and "
        "prioritise risk objectively. This project presents an Automated Information "
        "Security Risk Assessment Platform that lets an organisation enter its asset "
        "inventory and threat profile and then computes quantitative risk metrics "
        "following NIST SP 800-30 and the AssessITS methodology. The system calculates "
        "Single Loss Expectancy, Annualised Loss Expectancy, a composite risk score, "
        "and an AssessITS Risk Impact Rating that maps each asset-threat pair to a "
        "criticality band. A Flask backend exposes a REST API backed by SQLite, while "
        "a React dashboard visualises the prioritised risk register, a likelihood-"
        "versus-impact heat map, and a NIST 800-30 compliance checklist. The platform "
        "integrates the National Vulnerability Database to surface relevant CVEs and a "
        "large language model advisor to recommend security controls in plain English, "
        "each with graceful offline fallbacks. Downloadable PDF reports document the "
        "risk register, cost-benefit analysis, and compliance posture. Security testing "
        "confirms resilience to SQL injection and cross-site scripting, validated input "
        "handling, rate limiting, and stable behaviour under concurrent load, "
        "demonstrating a practical and reproducible quantitative risk tool."
    )
    add_body(doc, text)
    page_break(doc)


def toc(doc):
    """Section 3: Table of Contents with figure and table notes."""
    add_section_heading(doc, "Table of Contents", level=1)
    add_toc(doc)
    add_body(
        doc,
        "Note: The Table of Contents, List of Figures, and List of Tables are Word "
        "fields. Open the document in Microsoft Word and choose 'Update Field' "
        "(or press Ctrl+A then F9) to populate page numbers.",
        italic=True,
    )


def introduction(doc):
    """Section 4: Introduction."""
    add_section_heading(doc, "1. Introduction", level=1)
    add_section_heading(doc, "1.1 Background", level=2)
    add_body(
        doc,
        "Information security risk assessment is the process of identifying assets, "
        "the threats and vulnerabilities that endanger them, and the potential business "
        "impact of a successful attack. Established frameworks such as NIST SP 800-30 "
        "and ISO/IEC 27001 describe how to conduct such assessments, but applying them "
        "manually is time-consuming and inconsistent. Quantitative approaches express "
        "risk in monetary terms, enabling objective prioritisation and cost-benefit "
        "decisions about security controls.",
    )
    add_section_heading(doc, "1.2 Problem Statement", level=2)
    add_body(
        doc,
        "Many organisations, particularly small and medium enterprises, lack an "
        "accessible tool that combines quantitative risk formulas, current "
        "vulnerability intelligence, and clear control recommendations. Spreadsheets "
        "are error-prone and do not scale, while enterprise GRC suites are costly and "
        "complex. There is a need for a lightweight, reproducible platform that "
        "automates quantitative risk scoring and presents results that both technical "
        "staff and executives can act upon.",
    )
    add_section_heading(doc, "1.3 Objectives", level=2)
    add_numbered(
        doc,
        [
            "Implement quantitative risk formulas (SLE, ALE, composite risk score, and the AssessITS Risk Impact Rating) accurately and verifiably.",
            "Provide a REST API and database to manage assets, threats, and assessment results securely.",
            "Generate a prioritised risk register, cost-benefit analysis, and NIST 800-30 compliance checklist as downloadable PDF reports.",
            "Integrate CVE intelligence from the National Vulnerability Database and plain-English control recommendations via a large language model, with offline fallbacks.",
            "Deliver an interactive React dashboard for data entry and visualisation.",
            "Validate the system against common web security threats and concurrent load.",
        ],
    )
    add_section_heading(doc, "1.4 Scope and Limitations", level=2)
    add_body(
        doc,
        "The platform targets quantitative risk assessment for a single organisation "
        "session at a time and focuses on the calculation, reporting, and advisory "
        "workflow. It does not perform active network scanning or automated exploitation. "
        "CVE and LLM features depend on external services; when these are unavailable "
        "the system falls back to bundled data so that all functionality remains "
        "demonstrable offline. Asset valuation accuracy depends on user-supplied inputs.",
    )
    add_section_heading(doc, "1.5 Report Structure", level=2)
    add_body(
        doc,
        "Section 2 reviews related work and identifies the research gap. Section 3 "
        "describes the system design, including architecture, data flow, and threat "
        "model. Section 4 details the implementation and core algorithms. Section 5 "
        "presents the security analysis. Section 6 discusses results and limitations. "
        "Section 7 concludes with future work, followed by references and appendices.",
    )
    page_break(doc)


def literature_review(doc):
    """Section 5: Literature Review with 8+ works and a comparison table."""
    add_section_heading(doc, "2. Literature Review", level=1)
    add_body(
        doc,
        "Risk assessment methodologies fall broadly into qualitative, quantitative, and "
        "hybrid categories. This section reviews eight representative works and "
        "frameworks and identifies the gap this project addresses.",
    )
    add_bullets(
        doc,
        [
            "NIST SP 800-30 Rev. 1 [2] provides the canonical guide for conducting risk assessments, defining likelihood, impact, and the SLE/ALE quantitative model adopted here.",
            "Rahman et al.'s AssessITS [1] proposes a practical Risk Impact Rating combining asset value, threat value, and likelihood on a 1-250 band, which this platform implements directly.",
            "ISO/IEC 27001/27005 [3] specify an information security management system and risk treatment process, emphasising control selection and continual improvement.",
            "The FAIR model [4] offers a probabilistic factor analysis of information risk, expressing loss as distributions rather than point estimates.",
            "OCTAVE Allegro [5] is an asset-centric qualitative methodology suited to organisational self-assessment.",
            "The Common Vulnerability Scoring System (CVSS) [6] standardises severity scoring for vulnerabilities and underpins NVD data used in this platform.",
            "The MITRE ATT&CK knowledge base [7] catalogues adversary tactics and techniques, informing threat categorisation.",
            "Recent work on LLM-assisted security advisory [8] explores using large language models to translate technical findings into executive-readable guidance, mirroring this platform's advisor module.",
        ],
    )
    add_table_caption(doc, 1, "Comparison of risk assessment approaches.")
    add_table(
        doc,
        ["Approach", "Type", "Quantitative", "Automated", "CVE/LLM"],
        [
            ["NIST SP 800-30", "Framework", "Yes (SLE/ALE)", "Manual", "No"],
            ["AssessITS", "Method", "Yes (1-250)", "Partial", "No"],
            ["ISO 27005", "Framework", "Optional", "Manual", "No"],
            ["FAIR", "Model", "Yes (probabilistic)", "Tool-assisted", "No"],
            ["OCTAVE Allegro", "Method", "No (qualitative)", "Manual", "No"],
            ["This platform", "System", "Yes (SLE/ALE/AssessITS)", "Yes", "Yes"],
        ],
        col_widths=[1.5, 1.1, 1.5, 1.1, 1.0],
    )
    add_section_heading(doc, "2.1 Research Gap", level=2)
    add_body(
        doc,
        "While each reviewed work contributes a methodology or data source, none "
        "combines accurate quantitative formulas, live CVE intelligence, automated "
        "control recommendations, and executive-ready reporting in a single, "
        "reproducible, open tool. This platform fills that gap by integrating the "
        "AssessITS rating and NIST SP 800-30 quantitative model with NVD and LLM "
        "services behind a clean API and dashboard.",
    )
    page_break(doc)


def system_design(doc, figs):
    """Section 6: System Design with architecture, DFD, threat model, stack."""
    add_section_heading(doc, "3. System Design", level=1)
    add_section_heading(doc, "3.1 Architecture", level=2)
    add_body(
        doc,
        "The platform follows a classic three-tier architecture. A React single-page "
        "application communicates over HTTP with a Flask REST API, which persists data "
        "in SQLite and orchestrates the risk engine, report generator, CVE fetcher, and "
        "LLM advisor modules. Figure 1 shows the high-level architecture.",
    )
    add_figure(doc, figs["architecture"], 1, "System architecture of the platform.")
    add_section_heading(doc, "3.2 Component Descriptions", level=2)
    add_bullets(
        doc,
        [
            "Frontend: React 18 with Vite, Tailwind CSS, and Recharts for forms, dashboard, and charts.",
            "API layer: Flask blueprints exposing assets, threats, assessment, reports, compliance, and demo endpoints with a consistent {success, data, error} envelope.",
            "Risk engine: pure-Python module implementing all quantitative formulas.",
            "Persistence: SQLite accessed exclusively through parameterised queries.",
            "CVE fetcher: queries the NVD API v2.0 with a local mock fallback.",
            "LLM advisor: calls the Claude API with a rule-based JSON fallback.",
            "Report generator: ReportLab-based PDF export for register, CBA, and compliance.",
        ],
    )
    add_section_heading(doc, "3.3 Data-Flow Diagram", level=2)
    add_body(
        doc,
        "User input is validated and HTML-escaped, stored in SQLite, processed by the "
        "risk engine, enriched by the CVE and LLM modules, and finally rendered as a "
        "dashboard or exported as PDF. Figure 2 illustrates this flow.",
    )
    add_figure(doc, figs["dataflow"], 2, "Data-flow diagram.", width_inches=6.2)
    add_section_heading(doc, "3.4 Threat Model", level=2)
    add_body(
        doc,
        "A STRIDE analysis guided the security design. Each category is mapped to a "
        "concrete mitigating control implemented in the platform, as shown in Figure 3.",
    )
    add_figure(doc, figs["threat_model"], 3, "STRIDE threat model and mitigations.")
    add_section_heading(doc, "3.5 Technology Stack", level=2)
    add_table_caption(doc, 2, "Technology stack.")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Frontend", "React 18, Vite, Tailwind CSS, Recharts"],
            ["Backend", "Python 3.12, Flask, Flask-CORS"],
            ["Database", "SQLite (parameterised queries, WAL)"],
            ["Reporting", "ReportLab (PDF), python-docx (report)"],
            ["Intelligence", "NVD API v2.0, Anthropic Claude API"],
            ["Testing", "pytest"],
        ],
        col_widths=[1.8, 4.2],
    )
    page_break(doc)


def implementation(doc):
    """Section 7: Implementation with equations and code snippets."""
    add_section_heading(doc, "4. Implementation", level=1)
    add_section_heading(doc, "4.1 Development Approach", level=2)
    add_body(
        doc,
        "Development followed a backend-first, test-driven approach. The risk engine "
        "and its unit tests were built and verified against the NIST worked example "
        "before the API, frontend, and reporting layers were added. Work was committed "
        "incrementally to version control after each feature.",
    )
    add_section_heading(doc, "4.2 Core Risk Algorithms", level=2)
    add_body(
        doc,
        "The engine implements the following quantitative formulas. Single Loss "
        "Expectancy (SLE) is the expected loss from one incident:",
    )
    add_equation(doc, "SLE = AssetValue × ExposureFactor", 1)
    add_body(doc, "Annualised Loss Expectancy (ALE) scales SLE by the rate of occurrence:")
    add_equation(doc, "ALE = SLE × ARO", 2)
    add_body(doc, "The composite course risk score combines probability, vulnerability, mitigation, and uncertainty:")
    add_equation(doc, "R = (P × V) − M + U", 3)
    add_body(doc, "The AssessITS Risk Impact Rating maps each asset-threat pair to the 1-250 band:")
    add_equation(doc, "RiskImpactRating = AssetValue × ThreatValue × Likelihood", 4)
    add_body(
        doc,
        "where ThreatValue = ThreatLevel + VulnerabilityLevel and likelihood is mapped "
        "to a 0-5 scale. Listing 1 shows the corresponding implementation.",
    )
    add_code(
        doc,
        """def calculate_sle(asset_value, exposure_factor):
    if not 0 <= exposure_factor <= 1:
        raise ValueError("exposure_factor must be between 0 and 1")
    return asset_value * exposure_factor

def calculate_risk_impact_rating(asset_value, threat_value, likelihood):
    likelihood_scale = likelihood * 5  # map probability (0-1) to 0-5
    return asset_value * threat_value * likelihood_scale""",
    )
    add_section_heading(doc, "4.3 Risk Criticality Banding", level=2)
    add_body(
        doc,
        "Ratings are classified into criticality bands per AssessITS: 1-45 Low, "
        "46-99 Medium, 100-199 High, and 200-250 Critical. The register is sorted by "
        "impact rating so that the highest risks appear first.",
    )
    add_section_heading(doc, "4.4 Secure API Layer", level=2)
    add_body(
        doc,
        "All database access uses parameterised queries, and string inputs are length-"
        "validated and HTML-escaped before being returned. A per-IP rate limiter caps "
        "requests. Listing 2 shows the parameterised insert used for assets.",
    )
    add_code(
        doc,
        """conn.execute(
    "INSERT INTO assets (session_id, name, asset_type, value_usd, "
    "description, software) VALUES (?, ?, ?, ?, ?, ?)",
    (session_id, name, asset_type, value_usd, description, software),
)""",
    )
    add_section_heading(doc, "4.5 CVE and LLM Integration", level=2)
    add_body(
        doc,
        "The CVE fetcher queries the NVD API v2.0 using keyword search and respects the "
        "five-requests-per-thirty-seconds unauthenticated limit, falling back to bundled "
        "CVE data when offline. The LLM advisor sends the top risks to the Claude API "
        "with a strict-JSON system prompt and falls back to a rule-based control lookup "
        "keyed by criticality, ensuring the feature works without an API key.",
    )
    page_break(doc)


def security_analysis(doc):
    """Section 8: Security Analysis with attack scenarios and results table."""
    add_section_heading(doc, "5. Security Analysis", level=1)
    add_body(
        doc,
        "The platform was tested against common web application threats using an "
        "automated pytest suite. Table 3 summarises the attack scenarios and outcomes.",
    )
    add_table_caption(doc, 3, "Security test scenarios and results.")
    add_table(
        doc,
        ["Test", "Input / Method", "Expected", "Result"],
        [
            ["SQL Injection", "'; DROP TABLE assets; --", "Blocked by parameterised query", "Pass"],
            ["XSS", "<script>alert(1)</script>", "Output HTML-escaped", "Pass"],
            ["Input validation", "Over-length / invalid fields", "Rejected with 400", "Pass"],
            ["Rate limiting", "Burst > 100/min per IP", "429 after cap", "Pass"],
            ["LLM fallback", "No API key", "Rule-based recommendations", "Pass"],
            ["Concurrent load", "50 parallel assessments", "No crash, >=1 success", "Pass"],
        ],
        col_widths=[1.3, 1.9, 1.7, 0.8],
    )
    add_section_heading(doc, "5.1 Comparison with Baseline", level=2)
    add_body(
        doc,
        "Compared with a naive implementation that builds SQL strings and renders raw "
        "user input, the platform's parameterised queries and output escaping eliminate "
        "the injection and reflected-XSS vectors. Rate limiting adds a denial-of-service "
        "safeguard absent from a baseline Flask application.",
    )
    page_break(doc)


def results_discussion(doc, figs, register):
    """Section 9: Results and Discussion with charts and metrics."""
    add_section_heading(doc, "6. Results and Discussion", level=1)
    total_ale = sum(r["ale"] for r in register)
    crit_counts: dict[str, int] = {}
    for r in register:
        crit_counts[r["risk_criticality"]] = crit_counts.get(r["risk_criticality"], 0) + 1
    top = register[0]
    add_body(
        doc,
        f"Running the platform on a representative inventory of five assets and "
        f"{len(register)} asset-threat pairs produced a total Annualised Loss "
        f"Expectancy of approximately ${total_ale:,.0f}. The highest-rated risk was "
        f"'{top['threat_name']}' on '{top['asset_name']}' with an AssessITS impact "
        f"rating of {top['risk_impact_rating']:.0f} ({top['risk_criticality']}). "
        f"Figure 4 shows the distribution of risks across criticality bands, and "
        f"Figure 5 ranks the top risks by impact rating.",
    )
    add_figure(doc, figs["risk_distribution"], 4, "Risk distribution by criticality.", width_inches=4.5)
    add_figure(doc, figs["top_risks"], 5, "Top risks by AssessITS impact rating.", width_inches=5.8)
    add_figure(doc, figs["ale_by_asset"], 6, "Total Annualised Loss Expectancy by asset.", width_inches=5.8)
    add_table_caption(doc, 4, "Risk criticality distribution.")
    add_table(
        doc,
        ["Criticality", "Count"],
        [[k, str(crit_counts.get(k, 0))] for k in ["Critical", "High", "Medium", "Low"]],
        col_widths=[2.0, 1.5],
    )
    add_section_heading(doc, "6.1 Discussion and Limitations", level=2)
    add_body(
        doc,
        "The results demonstrate that the platform produces a coherent, prioritised "
        "risk picture from straightforward inputs, and that the AssessITS rating now "
        "stays within its intended 1-250 band after correcting an earlier scaling "
        "defect. Limitations include reliance on user-estimated asset values and "
        "probabilities, which introduce subjectivity, and dependence on external CVE "
        "and LLM services for live data. The composite score and AssessITS rating use "
        "different scales and are presented side by side rather than merged.",
    )
    page_break(doc)


def conclusion(doc):
    """Section 10: Conclusion."""
    add_section_heading(doc, "7. Conclusion", level=1)
    add_body(
        doc,
        "This project delivered a working Automated Information Security Risk "
        "Assessment Platform that meets its objectives: it implements verified "
        "quantitative risk formulas, exposes a secure API and database, generates "
        "prioritised reports and a compliance checklist, integrates CVE and LLM "
        "intelligence with offline fallbacks, and presents an interactive dashboard. "
        "Security testing confirmed resistance to injection and scripting attacks and "
        "stability under load.",
    )
    add_body(
        doc,
        "Key lessons learned include the value of test-driven development for "
        "numerical code—unit tests caught the impact-rating scaling defect—and the "
        "importance of graceful degradation so that external-service features remain "
        "demonstrable offline.",
    )
    add_section_heading(doc, "7.1 Future Work", level=2)
    add_bullets(
        doc,
        [
            "Add multi-user authentication and persistent organisation accounts.",
            "Incorporate probabilistic (Monte Carlo) loss distributions in addition to point estimates.",
            "Automate control-effectiveness tracking against the compliance checklist over time.",
            "Expand CVE matching using CPE identifiers rather than keyword search.",
            "Containerise the platform and add CI/CD for automated testing and deployment.",
        ],
    )
    page_break(doc)


def references(doc):
    """Section 11: IEEE references (8+, no Wikipedia)."""
    add_section_heading(doc, "References", level=1)
    refs = [
        "M. M. Rahman et al., \"AssessITS: Integrating procedural guidelines and practical evaluation metrics for organizational IT and cybersecurity risk assessment,\" arXiv:2410.01750, 2024.",
        "National Institute of Standards and Technology, \"Guide for Conducting Risk Assessments,\" NIST Special Publication 800-30 Rev. 1, 2012.",
        "International Organization for Standardization, \"ISO/IEC 27001:2022 Information security management systems — Requirements,\" 2022.",
        "J. Freund and J. Jones, Measuring and Managing Information Risk: A FAIR Approach. Butterworth-Heinemann, 2015.",
        "R. A. Caralli et al., \"Introducing OCTAVE Allegro: Improving the Information Security Risk Assessment Process,\" Carnegie Mellon SEI, CMU/SEI-2007-TR-012, 2007.",
        "FIRST.org, \"Common Vulnerability Scoring System v3.1: Specification Document,\" 2019.",
        "MITRE Corporation, \"MITRE ATT&CK: Design and Philosophy,\" MITRE, 2020.",
        "A. Happe and J. Cito, \"Getting pwn'd by AI: Penetration testing with large language models,\" in Proc. ACM ESEC/FSE, 2023.",
        "National Institute of Standards and Technology, \"National Vulnerability Database API v2.0,\" NIST, 2023.",
    ]
    for i, r in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"[{i}] {r}")
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    page_break(doc)


def appendices(doc):
    """Section 12: Appendices with GitHub link, user manual, raw test data."""
    add_section_heading(doc, "Appendices", level=1)
    add_section_heading(doc, "Appendix A: Source Code Repository", level=2)
    add_body(
        doc,
        f"The complete source code, including backend, frontend, tests, and this "
        f"report generator, is available on GitHub: {GITHUB_URL}. The commit history "
        f"documents incremental, feature-by-feature progress.",
    )
    add_section_heading(doc, "Appendix B: User Manual", level=2)
    add_numbered(
        doc,
        [
            "Start the backend: in backend/, activate the virtual environment and run 'python app.py'.",
            "Start the frontend: in frontend/, run 'npm install' then 'npm run dev' and open the shown URL.",
            "On the Assessment page, click 'Load Demo Data' or add assets and threats manually.",
            "Click 'Run Assessment' to compute the risk register.",
            "Open the Results page to view the dashboard, compliance checklist, and notifications.",
            "Use the Download Reports buttons to export the Risk Register, CBA, and Compliance PDFs.",
        ],
    )
    add_section_heading(doc, "Appendix C: Raw Test Data", level=2)
    add_body(
        doc,
        "The automated test suite (backend/tests) contains 19 tests across the risk "
        "engine and security scenarios; all pass. Representative results are summarised "
        "in Table 3. The seeded demo inventory used for the figures in this report "
        "comprises five assets and eight threats spanning data, software, hardware, and "
        "process asset types.",
    )


def build():
    """Generate figures and assemble the full .docx report."""
    figs = figmod.generate_all()
    register = figs["register"]

    doc = Document()
    set_base_styles(doc)

    # Front matter (no page numbers)
    cover_page(doc)
    abstract(doc)
    toc(doc)

    # Main body — start page numbering at 1 from the Introduction
    start_numbered_section(doc, restart_at=1)
    introduction(doc)
    literature_review(doc)
    system_design(doc, figs)
    implementation(doc)
    security_analysis(doc)
    results_discussion(doc, figs, register)
    conclusion(doc)
    references(doc)
    appendices(doc)

    doc.save(str(OUT_PATH))
    print(f"Report written to {OUT_PATH}")


if __name__ == "__main__":
    build()
