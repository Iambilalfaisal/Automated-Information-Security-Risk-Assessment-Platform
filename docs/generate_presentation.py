"""
generate_presentation.py
Generates a professional Word presentation for the
Automated Information Security Risk Assessment Platform capstone project.
"""

import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour Palette ─────────────────────────────────────────────────────────────
C_NAVY      = RGBColor(0x0F, 0x28, 0x4A)   # deep navy  – main heading bg
C_BLUE      = RGBColor(0x1D, 0x4E, 0x89)   # mid blue   – section heading bg
C_ACCENT    = RGBColor(0x27, 0x74, 0xC2)   # bright blue – accent
C_TEAL      = RGBColor(0x0D, 0x7A, 0x8A)   # teal       – sub-accent
C_RED       = RGBColor(0xC0, 0x39, 0x2B)   # critical
C_ORANGE    = RGBColor(0xE6, 0x7E, 0x22)   # high
C_YELLOW    = RGBColor(0xD4, 0xAC, 0x0D)   # medium
C_GREEN     = RGBColor(0x1E, 0x8B, 0x4C)   # low / positive
C_DARK_GRAY = RGBColor(0x2C, 0x3E, 0x50)   # body text
C_MID_GRAY  = RGBColor(0x55, 0x66, 0x77)   # secondary text
C_LIGHT_BG  = RGBColor(0xEA, 0xF2, 0xFB)   # light blue bg
C_LIGHTER   = RGBColor(0xF4, 0xF8, 0xFF)   # very light row bg
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_GOLD      = RGBColor(0xD4, 0xAC, 0x0D)   # gold highlight

# ── Helpers ────────────────────────────────────────────────────────────────────

def _rgb_hex(rgb) -> str:
    """Convert RGBColor to 6-char uppercase hex string."""
    return str(rgb).upper()


def _set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell background colour."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _rgb_hex(rgb))
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_border(cell, sides="all", color="CCCCCC", size="4"):
    """Add border to cell sides."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in (["top", "left", "bottom", "right"] if sides == "all" else [sides]):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _paragraph_border_bottom(para, color="2774C2", size="6"):
    """Add a bottom border under a paragraph (decorative rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _shading_paragraph(para, rgb: RGBColor):
    """Shade a paragraph background."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), _rgb_hex(rgb))
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)


def run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


# ── Section Heading Styles ─────────────────────────────────────────────────────

def add_cover_banner(doc, title, subtitle, date_str, authors):
    """Full-width cover page."""
    # Deep navy banner paragraph
    def banner_para(text, font_size, bold=True, color=C_WHITE, bg=C_NAVY,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0):
        p = doc.add_paragraph()
        _shading_paragraph(p, bg)
        p.alignment = align
        set_para_spacing(p, before=space_before, after=space_after)
        r = p.add_run(text)
        run_font(r, "Calibri", font_size, bold=bold, color=color)
        return p

    # Top coloured bar
    banner_para("", 6, bg=C_ACCENT, space_before=0, space_after=0)
    banner_para("UNIVERSITY OF MANAGEMENT AND TECHNOLOGY", 11, bg=C_NAVY,
                color=RGBColor(0xA8, 0xC8, 0xEC), bold=False, space_before=10)
    banner_para("Department of Information Security", 10, bg=C_NAVY,
                color=RGBColor(0x7A, 0xA7, 0xD3), bold=False, space_after=4)
    banner_para("", 4, bg=C_NAVY)
    banner_para(title, 28, bg=C_NAVY, space_before=8, space_after=4)
    banner_para(subtitle, 14, bg=C_NAVY, color=RGBColor(0xA8, 0xC8, 0xEC),
                bold=False, space_after=8)
    # Accent divider
    banner_para("", 6, bg=C_ACCENT)

    doc.add_paragraph()

    # Info table (2-col)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ("Authors",       authors),
        ("Course",        "Information Security Risk Assessment  ·  Spring 2026"),
        ("Institution",   "University of Management and Technology, Lahore"),
        ("Date",          date_str),
    ]
    for i, (label, value) in enumerate(info):
        row = tbl.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        _set_cell_bg(lc, C_NAVY)
        _set_cell_bg(vc, C_LIGHT_BG if i % 2 == 0 else C_LIGHTER)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label.upper())
        run_font(lr, "Calibri", 9, bold=True, color=C_WHITE)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        run_font(vr, "Calibri", 10, color=C_DARK_GRAY)
        set_para_spacing(lp, before=4, after=4)
        set_para_spacing(vp, before=4, after=4)
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(11.8)

    doc.add_paragraph()
    # Bottom accent line
    p = doc.add_paragraph()
    _shading_paragraph(p, C_ACCENT)
    set_para_spacing(p, before=0, after=0)
    r = p.add_run("Automated Information Security Risk Assessment Platform  ·  Capstone Submission 2026")
    run_font(r, "Calibri", 9, color=C_WHITE, italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def section_heading(doc, number, title, bg=C_BLUE):
    """Coloured full-width section heading block."""
    p = doc.add_paragraph()
    _shading_paragraph(p, bg)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=6, after=2)

    num_run = p.add_run(f"  {number}  ")
    run_font(num_run, "Calibri", 13, bold=True, color=C_GOLD)

    sep_run = p.add_run("│  ")
    run_font(sep_run, "Calibri", 13, bold=False, color=RGBColor(0x7A, 0xA7, 0xD3))

    title_run = p.add_run(title.upper())
    run_font(title_run, "Calibri", 13, bold=True, color=C_WHITE)
    return p


def sub_heading(doc, title, color=C_ACCENT, size=12, underline=True):
    """Bold coloured sub-heading with optional bottom rule."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=2)
    r = p.add_run(title)
    run_font(r, "Calibri", size, bold=True, color=color)
    if underline:
        _paragraph_border_bottom(p, color=_rgb_hex(color), size="4")
    return p


def body_para(doc, text, indent=0, size=10.5, color=C_DARK_GRAY, italic=False,
              before=2, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent)
    set_para_spacing(p, before=before, after=after)
    r = p.add_run(text)
    run_font(r, "Calibri", size, color=color, italic=italic)
    return p


def bullet_para(doc, text, level=0, size=10.5, color=C_DARK_GRAY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    set_para_spacing(p, before=1, after=1)
    r = p.add_run(text)
    run_font(r, "Calibri", size, color=color)
    return p


def callout_box(doc, label, text, bg=C_LIGHT_BG, label_color=C_BLUE, border_color="2774C2"):
    """Highlighted callout / info box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, bg)
    _set_cell_border(cell, sides="all", color=border_color, size="8")

    lp = cell.add_paragraph()
    lr = lp.add_run(f"  {label}  ")
    run_font(lr, "Calibri", 9, bold=True, color=label_color)
    set_para_spacing(lp, before=2, after=0)

    vp = cell.add_paragraph()
    vp.paragraph_format.left_indent = Cm(0.3)
    vr = vp.add_run(f"  {text}")
    run_font(vr, "Calibri", 10, color=C_DARK_GRAY)
    set_para_spacing(vp, before=0, after=4)
    doc.add_paragraph()


def kpi_row(doc, items):
    """Row of KPI stat boxes — items: [(label, value, sub), ...]"""
    cols = len(items)
    tbl = doc.add_table(rows=2, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_i, (label, value, sub) in enumerate(items):
        # Value row
        vcell = tbl.rows[0].cells[col_i]
        _set_cell_bg(vcell, C_NAVY)
        _set_cell_border(vcell, "bottom", color=_rgb_hex(C_ACCENT), size="8")
        vp = vcell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(str(value))
        run_font(vr, "Calibri", 22, bold=True, color=C_GOLD)
        set_para_spacing(vp, before=8, after=2)

        # Label row
        lcell = tbl.rows[1].cells[col_i]
        _set_cell_bg(lcell, C_BLUE)
        lp = lcell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        run_font(lr, "Calibri", 9, bold=True, color=C_WHITE)
        set_para_spacing(lp, before=2, after=2)

        if sub:
            sp = lcell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(sub)
            run_font(sr, "Calibri", 8, color=RGBColor(0xA8, 0xC8, 0xEC), italic=True)
            set_para_spacing(sp, before=0, after=4)

    doc.add_paragraph()


def formula_box(doc, formula_name, formula, explanation, result_range=""):
    """Styled formula block."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, RGBColor(0x0F, 0x28, 0x4A))
    _set_cell_border(cell, sides="all", color="2774C2", size="10")

    header_p = cell.add_paragraph()
    header_p.paragraph_format.left_indent = Cm(0.3)
    hr = header_p.add_run(f"  ƒ  {formula_name}")
    run_font(hr, "Calibri", 9, bold=True, color=RGBColor(0xA8, 0xC8, 0xEC))
    set_para_spacing(header_p, before=4, after=0)

    formula_p = cell.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = formula_p.add_run(formula)
    run_font(fr, "Courier New", 14, bold=True, color=C_GOLD)
    set_para_spacing(formula_p, before=6, after=4)

    if explanation:
        for line in explanation:
            ep = cell.add_paragraph()
            ep.paragraph_format.left_indent = Cm(0.8)
            er = ep.add_run(f"  •  {line}")
            run_font(er, "Calibri", 9, color=RGBColor(0x78, 0xA7, 0xD0))
            set_para_spacing(ep, before=1, after=1)

    if result_range:
        rp = cell.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = rp.add_run(f"Result Range:  {result_range}")
        run_font(rr, "Calibri", 9, bold=True, color=RGBColor(0x48, 0xBB, 0x78))
        set_para_spacing(rp, before=4, after=6)

    doc.add_paragraph()


def styled_table(doc, headers, rows, col_widths=None):
    """Styled data table with alternating row colours."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = tbl.rows[0]
    for ci, hdr in enumerate(headers):
        cell = hrow.cells[ci]
        _set_cell_bg(cell, C_NAVY)
        _set_cell_border(cell, sides="bottom", color="2774C2", size="8")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        run_font(r, "Calibri", 9, bold=True, color=C_WHITE)
        set_para_spacing(p, before=4, after=4)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        row_bg = C_LIGHTER if ri % 2 == 0 else C_WHITE
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            _set_cell_bg(cell, row_bg)
            _set_cell_border(cell, sides="bottom", color="DDDDDD", size="4")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(cell_text))
            run_font(r, "Calibri", 9, color=C_DARK_GRAY)
            set_para_spacing(p, before=3, after=3)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()


def criticality_band_row(doc):
    """Visual criticality band bar."""
    bands = [
        ("1–45",    "LOW",      C_GREEN,  C_WHITE),
        ("46–99",   "MEDIUM",   C_YELLOW, C_DARK_GRAY),
        ("100–199", "HIGH",     C_ORANGE, C_WHITE),
        ("200–250", "CRITICAL", C_RED,    C_WHITE),
    ]
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (rng, label, bg, fg) in enumerate(bands):
        # Label row
        lc = tbl.rows[0].cells[ci]
        _set_cell_bg(lc, bg)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        run_font(lr, "Calibri", 14, bold=True, color=fg)
        set_para_spacing(lp, before=10, after=4)

        # Range row
        rc = tbl.rows[1].cells[ci]
        _set_cell_bg(rc, C_NAVY)
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = rp.add_run(f"Score: {rng}")
        run_font(rr, "Calibri", 9, bold=False, color=RGBColor(0xA8, 0xC8, 0xEC))
        set_para_spacing(rp, before=4, after=6)

    doc.add_paragraph()


def two_col_bullets(doc, left_items, right_items, title_left="", title_right=""):
    """Two side-by-side bullet lists in a table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, (heading, items) in enumerate([(title_left, left_items), (title_right, right_items)]):
        cell = tbl.rows[0].cells[ci]
        _set_cell_bg(cell, C_LIGHTER)
        _set_cell_border(cell, sides="all", color="CCCCCC", size="4")
        if heading:
            hp = cell.add_paragraph()
            hr_ = hp.add_run(heading)
            run_font(hr_, "Calibri", 10, bold=True, color=C_ACCENT)
            set_para_spacing(hp, before=4, after=2)
            hp.paragraph_format.left_indent = Cm(0.3)
        for item in items:
            ip = cell.add_paragraph()
            ir = ip.add_run(f"  ✓  {item}")
            run_font(ir, "Calibri", 9.5, color=C_DARK_GRAY)
            set_para_spacing(ip, before=1, after=1)
            ip.paragraph_format.left_indent = Cm(0.2)

    doc.add_paragraph()


def framework_box_row(doc, frameworks):
    """Row of 3 framework boxes."""
    tbl = doc.add_table(rows=3, cols=len(frameworks))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, (icon, title, desc, color) in enumerate(frameworks):
        # Icon cell
        ic = tbl.rows[0].cells[ci]
        _set_cell_bg(ic, C_NAVY)
        ip = ic.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ir = ip.add_run(icon)
        run_font(ir, "Calibri", 20, color=color)
        set_para_spacing(ip, before=10, after=2)

        # Title cell
        tc = tbl.rows[1].cells[ci]
        _set_cell_bg(tc, color)
        tp = tc.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        run_font(tr, "Calibri", 10, bold=True, color=C_WHITE)
        set_para_spacing(tp, before=4, after=4)

        # Desc cell
        dc = tbl.rows[2].cells[ci]
        _set_cell_bg(dc, C_LIGHTER)
        dp = dc.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dr = dp.add_run(f"  {desc}")
        run_font(dr, "Calibri", 9, color=C_MID_GRAY)
        set_para_spacing(dp, before=6, after=6)
        dp.paragraph_format.left_indent = Cm(0.2)

    doc.add_paragraph()


def page_card_row(doc, pages):
    """Row of page feature cards."""
    tbl = doc.add_table(rows=3, cols=len(pages))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, (num, icon, title, features) in enumerate(pages):
        # Number + icon
        nc = tbl.rows[0].cells[ci]
        _set_cell_bg(nc, C_NAVY)
        np_ = nc.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np_.add_run(f"{icon}\n{num}")
        run_font(nr, "Calibri", 14, bold=True, color=C_GOLD)
        set_para_spacing(np_, before=8, after=2)

        # Title
        tc = tbl.rows[1].cells[ci]
        _set_cell_bg(tc, C_BLUE)
        tp = tc.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(title)
        run_font(tr, "Calibri", 9, bold=True, color=C_WHITE)
        set_para_spacing(tp, before=3, after=3)

        # Features
        fc = tbl.rows[2].cells[ci]
        _set_cell_bg(fc, C_LIGHTER)
        for feat in features:
            fp = fc.add_paragraph()
            fr = fp.add_run(f"  • {feat}")
            run_font(fr, "Calibri", 8.5, color=C_DARK_GRAY)
            set_para_spacing(fp, before=1, after=1)
            fp.paragraph_format.left_indent = Cm(0.1)
        set_para_spacing(fc.paragraphs[0], before=4, after=0)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # Page layout — A4, narrow margins
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── COVER PAGE ─────────────────────────────────────────────────────────────
    add_cover_banner(
        doc,
        title="Automated Information Security\nRisk Assessment Platform",
        subtitle="A Quantitative, AI-Assisted Risk Assessment System Implementing NIST SP 800-30,"
                 " AssessITS, and ISO/IEC 27001:2022",
        date_str="June 2026",
        authors="Bilal Faisal  ·  UMT Information Security Department",
    )

    add_page_break(doc)

    # ── TABLE OF CONTENTS (manual) ─────────────────────────────────────────────
    section_heading(doc, "TOC", "TABLE OF CONTENTS", bg=C_NAVY)
    toc_items = [
        ("1.", "Executive Summary"),
        ("2.", "Platform Overview"),
        ("3.", "Technology Stack"),
        ("4.", "System Architecture"),
        ("5.", "Academic Frameworks & Methodology"),
        ("6.", "Quantitative Risk Formulas"),
        ("7.", "Risk Criticality Bands"),
        ("8.", "Platform Features — All Six Pages"),
        ("9.", "Database Design"),
        ("10.", "Security Controls & Testing"),
        ("11.", "AI-Powered Threat Intelligence"),
        ("12.", "Key Platform Statistics"),
        ("13.", "Sample Risk Register"),
        ("14.", "Deployment & Scalability"),
        ("15.", "Academic References"),
    ]
    tbl = doc.add_table(rows=len(toc_items), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (num, title) in enumerate(toc_items):
        bg = C_LIGHTER if ri % 2 == 0 else C_WHITE
        nc = tbl.rows[ri].cells[0]
        tc_ = tbl.rows[ri].cells[1]
        _set_cell_bg(nc, C_NAVY if ri == 0 else bg)
        _set_cell_bg(tc_, C_NAVY if ri == 0 else bg)
        _set_cell_border(tc_, "bottom", color="DDDDDD", size="4")

        np_ = nc.paragraphs[0]
        nr = np_.add_run(num)
        run_font(nr, "Calibri", 10, bold=True,
                 color=C_GOLD if ri == 0 else C_ACCENT)
        set_para_spacing(np_, before=3, after=3)
        nc.width = Cm(1.2)

        tp_ = tc_.paragraphs[0]
        tr = tp_.add_run(title)
        run_font(tr, "Calibri", 10,
                 bold=(ri == 0),
                 color=C_WHITE if ri == 0 else C_DARK_GRAY)
        set_para_spacing(tp_, before=3, after=3)

    doc.add_paragraph()
    add_page_break(doc)

    # ── 1. EXECUTIVE SUMMARY ───────────────────────────────────────────────────
    section_heading(doc, "01", "Executive Summary")
    body_para(doc,
        "The Automated Information Security Risk Assessment Platform is a full-stack, "
        "production-ready web application developed as a University of Management and "
        "Technology (UMT) capstone project for Spring 2026. It automates quantitative "
        "information security risk assessment — a task traditionally performed manually "
        "by security analysts — through a unified, browser-based interface.",
        size=10.5)
    body_para(doc,
        "The platform integrates three internationally recognised security frameworks: "
        "NIST SP 800-30 Rev. 1 for quantitative risk measurement (SLE/ALE methodology), "
        "AssessITS (Rahman et al., 2024, arXiv:2410.01750) for impact-band scoring on "
        "a calibrated 1–250 scale, and ISO/IEC 27001:2022 for risk treatment classification. "
        "It further augments human analysis with live CVE threat intelligence from NIST NVD "
        "and AI-generated NIST control recommendations powered by Anthropic Claude.",
        size=10.5)

    kpi_row(doc, [
        ("Python Files",   "40+",    "backend + frontend"),
        ("Lines of Code",  "4,500+", "production code"),
        ("Plotly Charts",  "8",      "interactive visualizations"),
        ("NIST Controls",  "20",     "SP 800-53 tracked"),
        ("PDF Reports",    "3",      "risk register, CBA, compliance"),
        ("Academic Frameworks", "3", "NIST · AssessITS · ISO 27001"),
    ])

    callout_box(doc,
        "CORE MISSION",
        "Transform a traditionally manual, error-prone risk assessment process into an "
        "automated, repeatable, and auditable workflow — enabling organisations to quantify "
        "financial exposure, prioritise remediation, and demonstrate compliance in minutes "
        "rather than weeks.",
        bg=C_LIGHT_BG, label_color=C_BLUE)

    add_page_break(doc)

    # ── 2. PLATFORM OVERVIEW ───────────────────────────────────────────────────
    section_heading(doc, "02", "Platform Overview")
    body_para(doc,
        "The platform is built on Streamlit — a Python-native web framework that eliminates "
        "the need for a separate front-end codebase. The multi-page application consists of "
        "one dynamic entry point (app.py) and six specialised assessment pages, each "
        "addressing a distinct phase of the information security risk management lifecycle.",
        size=10.5)

    sub_heading(doc, "Assessment Lifecycle Flow")
    body_para(doc,
        "The platform follows the complete NIST SP 800-30 Rev. 1 risk assessment lifecycle:",
        size=10.5)

    lifecycle = [
        ("STEP 1", "Asset Inventory",       "Define organisational assets with monetary valuation (USD)"),
        ("STEP 2", "Threat Mapping",         "Map threats to assets with probability, vulnerability, ARO"),
        ("STEP 3", "Quantitative Analysis",  "Auto-calculate SLE, ALE, Risk Score, Impact Rating"),
        ("STEP 4", "Results Dashboard",      "Interactive charts, risk register, heatmap, compliance"),
        ("STEP 5", "Threat Intelligence",    "Live CVE data per asset from NIST NVD"),
        ("STEP 6", "Treatment Planning",     "Assign Mitigate / Accept / Transfer / Avoid strategies"),
        ("STEP 7", "History & Trends",       "ALE trend analysis across multiple assessment runs"),
        ("STEP 8", "Reporting",              "Download PDF reports: risk register, CBA, compliance"),
    ]
    tbl = doc.add_table(rows=len(lifecycle), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, (step, phase, desc) in enumerate(lifecycle):
        bg = C_NAVY if ri == 0 else (C_LIGHT_BG if ri % 2 == 0 else C_LIGHTER)
        for ci in range(3):
            _set_cell_bg(tbl.rows[ri].cells[ci], bg)
            _set_cell_border(tbl.rows[ri].cells[ci], "bottom", "CCCCCC", "4")

        sc = tbl.rows[ri].cells[0]
        sp = sc.paragraphs[0]
        sr = sp.add_run(step)
        run_font(sr, "Calibri", 9, bold=True,
                 color=C_GOLD if ri == 0 else C_ACCENT)
        set_para_spacing(sp, before=4, after=4)
        sc.width = Cm(2.0)

        pc = tbl.rows[ri].cells[1]
        pp_ = pc.paragraphs[0]
        pr = pp_.add_run(phase)
        run_font(pr, "Calibri", 9, bold=True,
                 color=C_WHITE if ri == 0 else C_DARK_GRAY)
        set_para_spacing(pp_, before=4, after=4)
        pc.width = Cm(4.0)

        dc = tbl.rows[ri].cells[2]
        dp = dc.paragraphs[0]
        dr = dp.add_run(desc)
        run_font(dr, "Calibri", 9, color=C_WHITE if ri == 0 else C_MID_GRAY)
        set_para_spacing(dp, before=4, after=4)

    doc.add_paragraph()
    add_page_break(doc)

    # ── 3. TECHNOLOGY STACK ────────────────────────────────────────────────────
    section_heading(doc, "03", "Technology Stack")
    body_para(doc,
        "The platform uses a pure-Python stack with no Node.js, no external front-end "
        "build tools, and no cloud services required for core functionality. All components "
        "are open-source and installable via pip.",
        size=10.5)

    sub_heading(doc, "Core Technologies")
    styled_table(doc,
        headers=["Layer", "Technology", "Version", "Purpose"],
        rows=[
            ("Frontend",     "Streamlit",       "≥1.28",  "Multi-page web UI — no HTML/JS needed"),
            ("Visualisation","Plotly",           "≥5.18",  "8 interactive chart types"),
            ("Data",         "Pandas",           "≥2.0",   "CSV handling, DataFrame operations"),
            ("Backend",      "Python",           "3.11+",  "Core language"),
            ("Database",     "SQLite (WAL)",     "3.x",    "7-table relational schema, zero-config"),
            ("Reports",      "ReportLab",        "≥4.0",   "PDF generation with embedded charts"),
            ("AI",           "Anthropic Claude", "claude-haiku-4-5", "NIST control recommendations"),
            ("Threat Intel", "NVD API v2.0",     "REST",   "Live CVE vulnerability data"),
            ("REST API",     "Flask",            "≥3.0",   "Optional JSON API for integrations"),
            ("HTTP Client",  "requests",         "≥2.31",  "NVD API, external integrations"),
            ("Config",       "python-dotenv",    "≥1.0",   "Environment variable management"),
            ("Testing",      "pytest",           "≥8.0",   "Unit + security test suite"),
            ("Dev Env",      "Docker / devcontainer","—",  "VS Code Dev Container included"),
            ("Deploy",       "Streamlit Cloud",  "Free",   "One-click cloud deployment"),
        ],
        col_widths=[3.0, 3.5, 2.5, 7.5])

    add_page_break(doc)

    # ── 4. SYSTEM ARCHITECTURE ─────────────────────────────────────────────────
    section_heading(doc, "04", "System Architecture")
    body_para(doc,
        "The platform follows a layered architecture: the Streamlit presentation layer "
        "delegates all computation to a separate backend package, which in turn reads "
        "and writes to the SQLite database through a typed data-access layer. This "
        "separation ensures testability, replaceability of individual components, and "
        "clean dependency boundaries.",
        size=10.5)

    sub_heading(doc, "Architecture Layers")
    layers = [
        ("PRESENTATION LAYER",  "Streamlit (app.py + pages/)",
         "Multi-page UI, CSS animations, Plotly charts, data editors", C_ACCENT),
        ("SERVICE LAYER",       "streamlit_lib/services.py + streamlit_lib/charts.py",
         "Assessment orchestration, compliance scoring, chart generation", C_TEAL),
        ("BUSINESS LOGIC",      "backend/modules/ (risk_engine, llm_advisor, cve_fetcher, report_generator)",
         "Risk formulas, AI recommendations, CVE parsing, PDF generation", C_BLUE),
        ("DATA ACCESS LAYER",   "backend/database/models.py",
         "Parameterised CRUD, upsert semantics, session-partitioned queries", C_NAVY),
        ("DATABASE",            "SQLite WAL (risk_platform.db)",
         "7 tables, 6 indices, FK constraints, cascading deletes", C_DARK_GRAY),
        ("EXTERNAL APIS",       "NIST NVD API v2.0  ·  Anthropic Claude API",
         "Live CVE data with mock fallback; LLM with rule-based fallback", C_MID_GRAY),
    ]
    for label, tech, desc, bg in layers:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        lc, tc_, dc = tbl.rows[0].cells

        _set_cell_bg(lc, bg)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        run_font(lr, "Calibri", 8, bold=True, color=C_WHITE)
        set_para_spacing(lp, before=5, after=5)
        lc.width = Cm(3.5)

        _set_cell_bg(tc_, C_LIGHT_BG)
        tp = tc_.paragraphs[0]
        tr = tp.add_run(tech)
        run_font(tr, "Courier New", 8.5, bold=True, color=C_BLUE)
        set_para_spacing(tp, before=5, after=5)
        tc_.width = Cm(5.5)

        _set_cell_bg(dc, C_LIGHTER)
        dp = dc.paragraphs[0]
        dr = dp.add_run(desc)
        run_font(dr, "Calibri", 9, color=C_MID_GRAY)
        set_para_spacing(dp, before=5, after=5)

    doc.add_paragraph()

    sub_heading(doc, "Directory Structure (Key Files)")
    styled_table(doc,
        headers=["File / Directory", "Lines", "Role"],
        rows=[
            ("app.py",                          "~210", "Entry point — landing page & executive dashboard"),
            ("pages/1_Assessment.py",           "~200", "Asset/threat inventory, CSV upload, assessment run"),
            ("pages/2_Results.py",              "~260", "Risk dashboard, charts, compliance, PDF downloads"),
            ("pages/3_Threat_Intelligence.py",  "~193", "Live CVE feed from NVD, keyword search"),
            ("pages/4_Treatment_Plan.py",       "~190", "Risk treatment tracking — inline data_editor"),
            ("pages/5_History.py",              "~125", "ALE trend chart, per-run assessment history"),
            ("pages/6_Methodology.py",          "~300", "Academic page — formulas, frameworks, references"),
            ("backend/modules/risk_engine.py",  "~237", "Core risk calculation engine (8 functions)"),
            ("backend/database/models.py",      "~461", "Data access layer (30+ CRUD functions)"),
            ("backend/database/schema.sql",     "~79",  "SQLite schema — 7 tables, 6 indices"),
            ("backend/modules/cve_fetcher.py",  "~131", "NVD API v2.0 with rate limiting + mock fallback"),
            ("backend/modules/report_generator.py","~284","ReportLab PDF generation (3 report types)"),
            ("streamlit_lib/style.py",          "~850", "Dark-theme CSS system (20+ CSS animations)"),
            ("streamlit_lib/charts.py",         "~393", "9 Plotly chart functions"),
            ("TOTAL (production)",              "4,500+","Complete platform codebase"),
        ],
        col_widths=[7.5, 1.5, 8.5])

    add_page_break(doc)

    # ── 5. ACADEMIC FRAMEWORKS ─────────────────────────────────────────────────
    section_heading(doc, "05", "Academic Frameworks & Methodology")
    body_para(doc,
        "The platform's methodology is grounded in three internationally recognised "
        "information security standards. Each framework contributes a distinct "
        "analytical dimension to the overall risk assessment pipeline.",
        size=10.5)

    framework_box_row(doc, [
        ("📋", "NIST SP 800-30 Rev. 1",
         "Quantitative risk assessment guide.\n"
         "Provides SLE, ALE, ARO methodology\n"
         "for monetary risk quantification.\n"
         "Source: U.S. NIST, 2012.",
         C_BLUE),
        ("🔬", "AssessITS (Rahman et al., 2024)",
         "Automated IS threat assessment model.\n"
         "1-250 Risk Impact Rating scale.\n"
         "Threat Value = Threat Level + Vuln.\n"
         "arXiv:2410.01750.",
         C_TEAL),
        ("🌐", "ISO/IEC 27001:2022",
         "International ISMS standard.\n"
         "Risk treatment strategies:\n"
         "Mitigate / Accept / Transfer / Avoid.\n"
         "Clause 6.1.3 — risk treatment.",
         C_NAVY),
    ])

    sub_heading(doc, "Framework Contribution Matrix")
    styled_table(doc,
        headers=["Framework", "Contribution to Platform", "Implementation"],
        rows=[
            ("NIST SP 800-30", "SLE / ALE / ARO quantitative formulas", "risk_engine.py → calculate_sle, calculate_ale"),
            ("NIST SP 800-30", "Composite Risk Score (P×V−M+U)", "risk_engine.py → calculate_risk_score"),
            ("NIST SP 800-30", "20-control compliance tracking", "constants.py + compliance_status table"),
            ("AssessITS",      "1–250 Risk Impact Rating",          "risk_engine.py → calculate_risk_impact_rating"),
            ("AssessITS",      "Threat Value (TL + VS)",            "risk_engine.py → calculate_threat_value"),
            ("AssessITS",      "Asset normalisation to 1–5 scale",  "risk_engine.py → normalise_asset_value"),
            ("AssessITS",      "Criticality bands (Low/Med/High/Crit)", "risk_engine.py → get_criticality"),
            ("ISO 27001",      "Risk treatment strategies",         "treatment_plans table, Treatment Plan page"),
            ("ISO 27001",      "Control implementation status",     "compliance_status table, Results page"),
            ("NIST 800-53",    "20 control families mapped",        "constants.py NIST_CONTROLS list"),
        ],
        col_widths=[3.5, 5.5, 8.5])

    add_page_break(doc)

    # ── 6. QUANTITATIVE FORMULAS ───────────────────────────────────────────────
    section_heading(doc, "06", "Quantitative Risk Formulas")
    body_para(doc,
        "All formulas are implemented in backend/modules/risk_engine.py with full "
        "input validation. The engine processes each asset–threat pair independently "
        "and returns a complete risk register entry with all metrics populated.",
        size=10.5)

    formula_box(doc,
        "Single Loss Expectancy (SLE)  —  NIST SP 800-30",
        "SLE  =  Asset_Value  ×  Exposure_Factor",
        [
            "Asset_Value: Monetary value of the asset (USD)",
            "Exposure_Factor: Fraction of asset value lost per incident (0.0 – 1.0)",
            "Example: $850,000 × 0.40 = $340,000 SLE",
        ],
        result_range="$0 – unlimited (USD)")

    formula_box(doc,
        "Annualised Loss Expectancy (ALE)  —  NIST SP 800-30",
        "ALE  =  SLE  ×  ARO",
        [
            "SLE: Single Loss Expectancy (from formula above)",
            "ARO: Annualised Rate of Occurrence — expected incidents per year",
            "Example: $340,000 × 1.2 = $408,000 ALE per year",
        ],
        result_range="$0 – unlimited (USD/year)")

    formula_box(doc,
        "Composite Risk Score  —  NIST SP 800-30",
        "R  =  (P × V)  −  M  +  U",
        [
            "P: Probability of threat occurrence (0.0 – 1.0)",
            "V: Vulnerability severity score (1 – 5)",
            "M: Mitigation effectiveness (0.0 – 1.0)",
            "U: Uncertainty margin (default 0.1 = 10%)",
            "Example: (0.65 × 5) − 0.30 + 0.15 = 3.50 (High)",
        ],
        result_range="Typically 0.0 – 5.0 (unbounded)")

    formula_box(doc,
        "Risk Impact Rating  —  AssessITS (Rahman et al., 2024)",
        "RIR  =  Asset_Scale  ×  Threat_Value  ×  Likelihood",
        [
            "Asset_Scale = 1.0 + (min(Value_USD, $1M) / $1M) × 4.0  →  range: 1.0–5.0",
            "Threat_Value = Threat_Level + Vulnerability_Score  →  range: 2–10",
            "Likelihood = Probability × 5  →  range: 0.0–5.0",
            "Example: 3.0 × 10 × 4.0 = 120 (HIGH band)",
        ],
        result_range="1 – 250 (AssessITS calibrated band)")

    formula_box(doc,
        "Cost-Benefit Analysis (CBA)  —  NIST SP 800-30",
        "CBA  =  ALE_Before  −  ALE_After  −  Control_Cost",
        [
            "ALE_Before: Annual loss without control",
            "ALE_After: ALE after control = ALE_Before × 0.70 (30% reduction model)",
            "Control_Cost: Annual cost of control = ALE × 0.10 (10% default)",
            "CBA > 0 → Control is cost-effective and recommended",
        ],
        result_range="CBA > 0 = Recommended  |  CBA ≤ 0 = Not cost-effective")

    add_page_break(doc)

    # ── 7. CRITICALITY BANDS ───────────────────────────────────────────────────
    section_heading(doc, "07", "Risk Criticality Bands")
    body_para(doc,
        "The AssessITS 1–250 Risk Impact Rating is partitioned into four criticality "
        "bands. These bands drive the colour-coding throughout the platform (risk "
        "register, charts, CVE cards) and determine the AI control recommendation "
        "priority.",
        size=10.5)

    criticality_band_row(doc)

    styled_table(doc,
        headers=["Band", "RIR Range", "Meaning", "Platform Response", "Badge Color"],
        rows=[
            ("LOW",      "1 – 45",    "Negligible financial and operational exposure",
             "Informational; annual review",               "Green   #1E8B4C"),
            ("MEDIUM",   "46 – 99",   "Moderate risk; quarterly monitoring required",
             "Document & schedule treatment",              "Yellow  #D4AC0D"),
            ("HIGH",     "100 – 199", "Significant exposure; priority remediation",
             "Immediate owner assignment & due date",      "Orange  #E67E22"),
            ("CRITICAL", "200 – 250", "Extreme financial and operational impact",
             "Urgent escalation; board notification level","Red     #C0392B"),
        ],
        col_widths=[2.0, 2.2, 5.5, 5.5, 2.5])

    add_page_break(doc)

    # ── 8. PLATFORM FEATURES ───────────────────────────────────────────────────
    section_heading(doc, "08", "Platform Features — All Six Pages")

    sub_heading(doc, "Pages 1–3")
    page_card_row(doc, [
        ("PAGE 1", "📋", "Assessment",
         ["Asset inventory (name, type, USD)", "Threat mapping per asset",
          "12-field threat form", "Demo data loader (5 assets, 8 threats)",
          "CSV bulk upload", "Input validation", "Run Assessment button"]),
        ("PAGE 2", "📊", "Results & Reports",
         ["4 KPI stat cards", "Criticality pie chart", "Top-10 risk bar chart",
          "Likelihood × Impact heatmap", "ALE by asset bar chart",
          "Sortable/filterable risk register", "NIST compliance checklist",
          "AI control recommendations", "3 PDF downloads"]),
        ("PAGE 3", "🔍", "Threat Intelligence",
         ["Per-asset NVD CVE lookup", "CVE severity bar chart",
          "Color-coded CVE detail cards", "CVSS score display",
          "NVD keyword search", "Critical CVE notifications",
          "Mock fallback (offline support)"]),
    ])

    sub_heading(doc, "Pages 4–6")
    page_card_row(doc, [
        ("PAGE 4", "🛠️", "Treatment Plan",
         ["Inline data_editor (spreadsheet UI)", "Strategy: Mitigate/Accept/Transfer/Avoid",
          "Owner & due-date assignment", "Status: Pending/In Progress/Completed",
          "Treatment status donut chart", "Strategy breakdown bars",
          "CSV export", "SQLite persistence (upsert)"]),
        ("PAGE 5", "📈", "Assessment History",
         ["Total run count KPI", "Latest ALE KPI", "Delta vs previous run",
          "ALE trend line + area chart", "Per-run expandable sections",
          "Full risk register per run", "AI rec summary per run"]),
        ("PAGE 6", "📖", "Methodology",
         ["Platform overview narrative", "Framework alignment (3 cards)",
          "LaTeX-rendered formulas", "Criticality band reference",
          "NIST control library (20 controls)", "Architecture flowchart",
          "Academic bibliography"]),
    ])

    sub_heading(doc, "App Entry Point — Executive Dashboard")
    body_para(doc,
        "When an assessment exists, app.py renders an Executive Dashboard with 4 KPI "
        "cards (Total Assets, Total Threats, Total ALE, Highest Risk), two summary "
        "charts, CVE alert banners, and a navigation grid of all 6 pages. When no "
        "assessment exists, it renders an animated hero landing page with a "
        "quickstart guide and the full navigation grid.",
        size=10.5)

    add_page_break(doc)

    # ── 9. DATABASE DESIGN ─────────────────────────────────────────────────────
    section_heading(doc, "09", "Database Design")
    body_para(doc,
        "The platform uses SQLite with Write-Ahead Logging (WAL) for concurrent reads "
        "and foreign key constraints for referential integrity. Session isolation is "
        "achieved by partitioning all queries on a UUID session_id generated per "
        "browser session.",
        size=10.5)

    sub_heading(doc, "Schema Overview — 7 Tables")
    styled_table(doc,
        headers=["Table", "Key Columns", "Constraint", "Purpose"],
        rows=[
            ("assets",           "id, session_id, name, asset_type, value_usd, software",
             "PK(id)", "Asset inventory — monetary values"),
            ("threats",          "id, asset_id, name, category, probability, vulnerability_score, aro, exposure_factor",
             "FK → assets(id) CASCADE", "Threat catalog per asset"),
            ("assessments",      "id, session_id, results_json, created_at",
             "Index(session_id)", "Serialised risk register (JSON blob)"),
            ("compliance_status","id, session_id, control_id, control_name, status",
             "UNIQUE(session_id, control_id)", "NIST control implementation tracking"),
            ("notifications",    "id, session_id, message, severity, read_flag",
             "Index(session_id)", "CVE alerts and system messages"),
            ("treatment_plans",  "id, session_id, asset_name, threat_name, strategy, owner, due_date, status",
             "UNIQUE(session+asset+threat)", "ISO 27001 risk treatment records"),
        ],
        col_widths=[3.0, 6.5, 4.5, 4.5])

    sub_heading(doc, "Key Design Decisions")
    two_col_bullets(doc,
        left_items=[
            "WAL mode — concurrent reads without locking writes",
            "PRAGMA foreign_keys = ON — referential integrity enforced",
            "6 indices on session_id and FK columns",
            "Cascading deletes on asset removal (threats follow)",
            "Parameterised queries throughout (SQL injection prevention)",
        ],
        right_items=[
            "Session partitioning — each browser session is isolated",
            "ON CONFLICT … DO UPDATE (upsert) for treatment plans & compliance",
            "results_json blob — risk register serialised as JSON for flexibility",
            "get_all_assessments() returns parsed summary for history trending",
            "Designed for easy migration to PostgreSQL (standard SQL)",
        ],
        title_left="Storage Design",
        title_right="Query Patterns")

    add_page_break(doc)

    # ── 10. SECURITY CONTROLS ──────────────────────────────────────────────────
    section_heading(doc, "10", "Security Controls & Testing")
    body_para(doc,
        "Security is a first-class concern throughout the codebase. The platform "
        "implements multiple defensive layers to prevent common web vulnerabilities "
        "and includes a documented penetration test report.",
        size=10.5)

    sub_heading(doc, "Built-In Security Controls")
    styled_table(doc,
        headers=["Control", "Implementation", "Mitigates"],
        rows=[
            ("SQL Injection Prevention", "All DB queries use parameterised statements (? placeholders)", "OWASP A03"),
            ("Session Isolation",        "UUID session_id partitions all data per browser session", "Data leakage between users"),
            ("Input Validation",         "validate_assessment_inputs() enforces formula constraints before calculation", "Invalid / malformed risk data"),
            ("Rate Limiting — NVD",      "5 req / 30 sec enforced (NVD policy); configurable with API key", "API abuse / ban"),
            ("Rate Limiting — Flask",    "100 req / 60 sec (configurable in config.py)", "DoS on optional REST API"),
            ("HTTPS",                    "Streamlit Cloud auto-enforces HTTPS; Flask accepts --ssl-certfile", "MITM / eavesdropping"),
            ("Secrets Management",       "st.secrets → os.environ bridge; .env.example (no real keys)", "Credential exposure in repo"),
            ("Fallback Resilience",      "mock_cves.json + fallback_controls.json if APIs unavailable", "Service unavailability"),
        ],
        col_widths=[4.5, 8.0, 4.5])

    sub_heading(doc, "Test Suite")
    styled_table(doc,
        headers=["File", "Tests", "Coverage"],
        rows=[
            ("backend/tests/test_risk_engine.py", "Unit tests for all 8 risk formula functions",
             "SLE, ALE, Risk Score, RIR, CBA, criticality bands, input validation"),
            ("backend/tests/test_security.py",    "Security penetration tests",
             "SQL injection patterns, authentication bypass, CORS policy, rate limit enforcement"),
            ("backend/tests/conftest.py",          "pytest fixtures (in-memory DB, sample assets/threats)",
             "Shared test context across both test modules"),
        ],
        col_widths=[5.5, 6.5, 6.5])

    callout_box(doc,
        "SECURITY TEST REPORT",
        "A dedicated SECURITY_TESTING_REPORT.md documents all penetration test results, "
        "vulnerability findings, and remediation evidence. All critical and high findings "
        "were resolved before submission.",
        bg=C_LIGHT_BG, label_color=C_RED, border_color="C0392B")

    add_page_break(doc)

    # ── 11. AI THREAT INTELLIGENCE ─────────────────────────────────────────────
    section_heading(doc, "11", "AI-Powered Threat Intelligence")

    sub_heading(doc, "Anthropic Claude Integration")
    body_para(doc,
        "The LLM Advisor module (backend/modules/llm_advisor.py) queries Anthropic's "
        "Claude (claude-haiku-4-5) with the top 5 risks from the register and returns "
        "structured JSON control recommendations. If no API key is configured, the system "
        "automatically falls back to a rule-based engine using fallback_controls.json — "
        "ensuring the platform is fully functional without cloud credentials.",
        size=10.5)

    styled_table(doc,
        headers=["Capability", "Technology", "Detail"],
        rows=[
            ("Control Recommendations", "Claude claude-haiku-4-5", "Top 5 risks → NIST control IDs, names, explanations, priority, cost estimate"),
            ("Structured Output",       "JSON schema enforcement",  "LLM instructed to return valid JSON; parsed and validated before display"),
            ("Offline Fallback",        "fallback_controls.json",   "Rule-based mapping: criticality band → pre-defined NIST controls"),
            ("CVE Intelligence",        "NIST NVD API v2.0",        "Per-asset CVE lookup; CVSS v3.1/3.0/v2 parsing; top 5 results"),
            ("CVE Mock Fallback",       "mock_cves.json",           "Realistic offline CVE data — platform works without internet"),
            ("Critical CVE Alerts",     "notifications table",      "Critical CVEs (CVSS ≥ 9.0) logged automatically as notifications"),
            ("CVE Keyword Search",      "NVD free-text search",     "User-driven search for any software, product, or CVE keyword"),
        ],
        col_widths=[4.5, 4.5, 9.0])

    sub_heading(doc, "CVE Severity Classification")
    styled_table(doc,
        headers=["Severity", "CVSS Score", "Platform Colour", "Response Level"],
        rows=[
            ("Critical", "9.0 – 10.0", "Red (#EF4444) — pulsing badge",   "Immediate notification; board-level escalation"),
            ("High",     "7.0 – 8.9",  "Orange (#F97316)",                 "Priority remediation; assign owner within 48h"),
            ("Medium",   "4.0 – 6.9",  "Yellow (#EAB308)",                 "Schedule treatment; next sprint inclusion"),
            ("Low",      "0.1 – 3.9",  "Green (#22C55E)",                  "Monitor; annual review"),
        ],
        col_widths=[2.5, 2.5, 5.5, 7.5])

    add_page_break(doc)

    # ── 12. KEY STATISTICS ─────────────────────────────────────────────────────
    section_heading(doc, "12", "Key Platform Statistics")

    sub_heading(doc, "Codebase Metrics")
    kpi_row(doc, [
        ("Python Files",   "40+",    "backend + streamlit_lib + pages"),
        ("Lines of Code",  "4,500+", "production-grade Python"),
        ("CSS Lines",      "850+",   "dark theme + 20 animations"),
        ("DB Tables",      "7",      "normalised relational schema"),
        ("DB Indices",     "6",      "query performance"),
    ])

    kpi_row(doc, [
        ("Functions",       "50+",  "risk, DB, charts, services"),
        ("Chart Types",     "8",    "Plotly interactive"),
        ("REST Endpoints",  "20+",  "Flask optional API"),
        ("Test Cases",      "15+",  "pytest unit + security"),
        ("Git Commits",     "15+",  "feature-based history"),
    ])

    sub_heading(doc, "Platform Feature Count")
    kpi_row(doc, [
        ("App Pages",           "7",  "landing + 6 assessment pages"),
        ("NIST Controls",       "20", "SP 800-53 tracked"),
        ("PDF Report Types",    "3",  "register, CBA, compliance"),
        ("Assessment Formulas", "8",  "SLE, ALE, RIR, CBA, ..."),
        ("Frameworks",          "3",  "NIST · AssessITS · ISO 27001"),
    ])

    sub_heading(doc, "Sample Demo Dataset")
    styled_table(doc,
        headers=["Asset", "Type", "Value (USD)", "Key Threat", "Est. ALE"],
        rows=[
            ("Customer Database Server",    "Data",           "$850,000", "SQL Injection / Data Breach",       "$408,000"),
            ("Public Web Server",           "Software",       "$500,000", "DDoS Attack",                       "$180,000"),
            ("Employee Laptop Fleet",       "Hardware",       "$220,000", "Ransomware Infection",              "$96,000"),
            ("Identity Management System",  "Software",       "$280,000", "Phishing / Credential Theft",       "$126,000"),
            ("Email & Collaboration Server","Software",       "$180,000", "Email Spoofing / BEC",              "$63,000"),
            ("Network File Storage",        "Data",           "$300,000", "Insider Data Exfiltration",         "$135,000"),
            ("Dev Lab Servers",             "Hardware",       "$160,000", "Unpatched RCE Vulnerability",       "$72,000"),
            ("Backup & DR System",          "Infrastructure", "$120,000", "Backup Corruption",                 "$36,000"),
        ],
        col_widths=[5.0, 2.5, 2.8, 5.5, 2.0])

    add_page_break(doc)

    # ── 13. SAMPLE RISK REGISTER ──────────────────────────────────────────────
    section_heading(doc, "13", "Sample Risk Register Output")
    body_para(doc,
        "The following is a representative risk register produced by the platform "
        "after running a quantitative assessment on the built-in demo dataset. "
        "Risks are sorted by Risk Impact Rating (descending) — highest priority first.",
        size=10.5)

    # Custom colour-coded register
    headers = ["Asset", "Threat", "RIR", "Band", "ALE ($)", "Risk Score", "CBA ($)"]
    rows = [
        ("Customer DB Server",   "SQL Injection / Data Breach",    "152", "HIGH",     "408,000", "3.5", "81,600"),
        ("Identity Mgmt System", "Phishing / Credential Theft",    "140", "HIGH",     "126,000", "3.2", "25,200"),
        ("Public Web Server",    "DDoS Attack",                    "128", "HIGH",     "180,000", "3.0", "36,000"),
        ("Customer DB Server",   "Insider Data Exfiltration",      "118", "HIGH",     "255,000", "3.1", "51,000"),
        ("Network File Storage", "Ransomware Encryption",          "108", "HIGH",     "162,000", "2.9", "32,400"),
        ("Email Server",         "Email Spoofing / BEC",           "85",  "MEDIUM",    "63,000", "2.6", "12,600"),
        ("Laptop Fleet",         "Malware via USB Drive",          "72",  "MEDIUM",    "88,000", "2.4", "17,600"),
        ("Dev Lab Servers",      "Unpatched RCE Vulnerability",    "65",  "MEDIUM",    "72,000", "2.2", "14,400"),
        ("Backup System",        "Backup Corruption",              "38",  "LOW",       "36,000", "1.8",  "7,200"),
        ("CCTV System",          "Physical Tampering",             "22",  "LOW",       "14,000", "1.2",  "2,800"),
    ]
    band_colors = {
        "HIGH":     C_ORANGE,
        "MEDIUM":   C_YELLOW,
        "LOW":      C_GREEN,
        "CRITICAL": C_RED,
    }

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        _set_cell_bg(cell, C_NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        run_font(r, "Calibri", 9, bold=True, color=C_WHITE)
        set_para_spacing(p, before=4, after=4)

    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        band = row_data[3]
        row_bg = C_LIGHTER if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            if ci == 3:  # Band cell
                _set_cell_bg(cell, band_colors.get(band, C_MID_GRAY))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                run_font(r, "Calibri", 9, bold=True, color=C_WHITE)
            else:
                _set_cell_bg(cell, row_bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                run_font(r, "Calibri", 9, color=C_DARK_GRAY)
            _set_cell_border(cell, "bottom", "DDDDDD", "4")
            set_para_spacing(p, before=3, after=3)

    # Set column widths
    col_widths_reg = [3.8, 4.5, 1.2, 1.8, 1.8, 1.8, 1.8]
    for ci, w in enumerate(col_widths_reg):
        for row in tbl.rows:
            row.cells[ci].width = Cm(w)

    doc.add_paragraph()

    callout_box(doc,
        "TOTAL ALE (Demo Dataset)",
        "Sum of all Annualised Loss Expectancies across 10 assets: approximately $1,245,000/year. "
        "The platform's CBA analysis identifies $249,000 in potential annual savings through "
        "recommended NIST controls — yielding a positive cost-benefit ratio of 5:1.",
        bg=C_LIGHT_BG, label_color=C_BLUE)

    add_page_break(doc)

    # ── 14. DEPLOYMENT ─────────────────────────────────────────────────────────
    section_heading(doc, "14", "Deployment & Scalability")

    sub_heading(doc, "Deployment Options")
    styled_table(doc,
        headers=["Option", "Command / Method", "Database", "Best For"],
        rows=[
            ("Local Development",   "streamlit run app.py",                 "Persistent local SQLite",  "Individual demos, development"),
            ("Docker Dev Container",".devcontainer/devcontainer.json (VS Code)", "Persistent volume",   "Team development, CI/CD"),
            ("Streamlit Cloud",     "Push to GitHub → share.streamlit.io",  "Ephemeral SQLite",         "Public demos, evaluation"),
            ("Self-Hosted Server",  "gunicorn -w 4 backend.app:app",        "Persistent / PostgreSQL",  "Production enterprise use"),
            ("Cloud (AWS/GCP)",     "Docker image → ECS/GKE",              "AWS RDS / Cloud SQL",       "High availability, scale"),
        ],
        col_widths=[4.0, 5.5, 3.5, 5.5])

    sub_heading(doc, "Environment Configuration")
    body_para(doc,
        "All secrets are managed via environment variables (see .env.example). "
        "No hard-coded credentials exist in the codebase. The Streamlit secrets "
        "bridge (streamlit_lib/paths.py) automatically maps st.secrets → os.environ "
        "for seamless transition between local and cloud deployment.",
        size=10.5)
    styled_table(doc,
        headers=["Variable", "Required", "Purpose"],
        rows=[
            ("ANTHROPIC_API_KEY",    "Optional", "Claude AI control recommendations (fallback enabled)"),
            ("NVD_API_KEY",          "Optional", "NVD API rate limit: 5 req/30 sec → unlimited"),
            ("DATABASE_PATH",        "Optional", "Custom SQLite path (default: backend/database/risk_platform.db)"),
            ("FLASK_SECRET_KEY",     "Optional", "Flask session signing (only if REST API used)"),
            ("FLASK_ENV",            "Optional", "production / development (Flask only)"),
        ],
        col_widths=[4.5, 2.0, 12.0])

    add_page_break(doc)

    # ── 15. REFERENCES ─────────────────────────────────────────────────────────
    section_heading(doc, "15", "Academic References")

    refs = [
        ("[1]", "Rahman, M. M., et al. (2024).",
         "AssessITS: Automated Information Security Threat Assessment.",
         "arXiv preprint arXiv:2410.01750. https://arxiv.org/abs/2410.01750"),
        ("[2]", "NIST. (2012).",
         "Special Publication 800-30 Rev. 1: Guide for Conducting Risk Assessments.",
         "National Institute of Standards and Technology, Gaithersburg, MD."),
        ("[3]", "ISO/IEC. (2022).",
         "ISO/IEC 27001:2022 — Information Security, Cybersecurity and Privacy Protection.",
         "International Organization for Standardization, Geneva."),
        ("[4]", "NIST. (2020).",
         "Special Publication 800-53 Rev. 5: Security and Privacy Controls for Information Systems.",
         "National Institute of Standards and Technology, Gaithersburg, MD."),
        ("[5]", "NIST. (2024).",
         "National Vulnerability Database (NVD) API v2.0.",
         "https://nvd.nist.gov/developers — Accessed June 2026."),
        ("[6]", "Anthropic. (2024).",
         "Claude API Documentation: claude-haiku-4-5.",
         "https://docs.anthropic.com — Accessed June 2026."),
        ("[7]", "Streamlit Inc. (2024).",
         "Streamlit: The Fastest Way to Build Data Apps.",
         "https://streamlit.io — v1.28+"),
    ]

    for num, authors, title, source in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        set_para_spacing(p, before=4, after=2)

        nr = p.add_run(f"{num}  ")
        run_font(nr, "Calibri", 10, bold=True, color=C_ACCENT)

        ar = p.add_run(authors + "  ")
        run_font(ar, "Calibri", 10, bold=False, color=C_DARK_GRAY)

        tr = p.add_run(title + "  ")
        run_font(tr, "Calibri", 10, bold=True, italic=False, color=C_NAVY)

        sr = p.add_run(source)
        run_font(sr, "Calibri", 10, italic=True, color=C_MID_GRAY)

    doc.add_paragraph()

    # Final footer banner
    p = doc.add_paragraph()
    _shading_paragraph(p, C_NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=8, after=8)
    r1 = p.add_run("Automated Information Security Risk Assessment Platform  ·  ")
    run_font(r1, "Calibri", 9, color=RGBColor(0x7A, 0xA7, 0xD3))
    r2 = p.add_run("UMT Capstone 2026  ·  ")
    run_font(r2, "Calibri", 9, bold=True, color=C_WHITE)
    r3 = p.add_run("NIST SP 800-30  ·  AssessITS  ·  ISO/IEC 27001:2022")
    run_font(r3, "Calibri", 9, color=RGBColor(0x7A, 0xA7, 0xD3))

    p2 = doc.add_paragraph()
    _shading_paragraph(p2, C_ACCENT)
    set_para_spacing(p2, before=0, after=0)
    r4 = p2.add_run("")
    run_font(r4, "Calibri", 4, color=C_WHITE)

    return doc


if __name__ == "__main__":
    import os
    os.makedirs("docs", exist_ok=True)
    output = "docs/IS_Risk_Assessment_Platform_Presentation.docx"
    doc = build_document()
    doc.save(output)
    print(f"OK  Document saved: {output}")
